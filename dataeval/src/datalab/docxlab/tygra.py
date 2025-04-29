import warnings
import collections


from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Mm

"""
    Office xml codes are from: http://officeopenxml.com/index.php
    python-docx source: https://python-docx.readthedocs.io/en/latest/index.html
"""


class Spacer:
    def __init__(self, *args, **kwargs):
        return

    @staticmethod
    def add_to_document(document):
        document.add_paragraph('', style='CNormal')


class PageBreak:
    def __init__(self, *args, **kwargs):
        return

    @staticmethod
    def add_to_document(document):
        document.add_page_break()


class Text(object):
    def __init__(self, text, **kwargs):
        """
        Basic element of the document, any raw string stored in these has the same style,
            which are stored in the properties: bold, italic, underline, size.
        If strings should have different style in one paragraph then TextCollection should be used,
            which stores Text in a tuple.

        :param text:
            Create text that can be formatted.
        :param kwargs:
            bold : bool
                Bold text.
            italic : bool
                Italic text.
            underline : bool
                Underlined text.
            size : int
                Size of text
        """
        # Check if it is a recreation
        if isinstance(text, Text):
            # get inherited style

            self._text = text._text
            self._bold = text._bold
            self._italic = text._italic
            self._underline = text._underline
            self._size = text._size

            # get overrides only if custom formatting was not assigned
            if text._text is not None:
                self._text = text._text
            if self._bold is None:
                self._bold = kwargs.pop('bold', self._bold)
            if self._underline is None:
                self._underline = kwargs.pop('underline', self._underline)
            if self._italic is None:
                self._italic = kwargs.pop('italic', self._italic)
            if self._size == 'default':
                self._size = kwargs.pop('size', self._size)
            return
        # It is a raw string so init everything
        elif isinstance(text, basestring):
            self._bold = kwargs.pop('bold', None)
            self._underline = kwargs.pop('underline', None)
            self._italic = kwargs.pop('italic', None)
            self._text = text
            self._size = kwargs.pop('size', 'default')
            return
        else:
            raise TypeError("The text argument should be an instance of 'Text' or 'basestring',"
                            " not '{}'".format(type(text)))

    def __str__(self):
        warnings.warn("Format will be lost! Use 'add_to_paragraph'")
        return self._text

    def __repr__(self):
        warnings.warn("Format will be lost! Use 'add_to_paragraph'")
        return self._text

    def __add__(self, other):
        # Adding it to an other 'basestring', 'Text' or 'TextCollection' (or to a tuple of these)
        # A new TextCollection will be made in order to keep the different styles
        return TextCollection([self, other])

    def __radd__(self, other):
        # Adding it to an other 'basestring', 'Text' or 'TextCollection' (or to a tuple of these)
        # A new TextCollection will be made in order to keep the different styles
        return TextCollection([other, self])

    def add_to_paragraph(self, paragraph):
        # Append to the paragraph with 'add_run'
        r = paragraph.add_run(self._text)
        # Add stored styles
        if self._size == 'default':
            r.font.size = Pt(12)  # create the proper object from int
        else:
            r.font.size = Pt(self._size)  # create the proper object from int
        r.bold = self._bold
        r.underline = self._underline
        r.italic = self._italic

    @property
    def text(self):
        warnings.warn("Format will be lost! Use 'add_to_paragraph'")
        return self._text

    @text.setter
    def text(self, text, **kwargs):
        self._bold = kwargs.pop('bold', False)
        self._underline = kwargs.pop('underline', False)
        self._italic = kwargs.pop('italic', False)
        self._size = kwargs.pop('size', 'default')
        self._text = text

    @property
    def bold(self):
        return self._bold

    @bold.setter
    def bold(self, is_bold):
        self._bold = is_bold

    @property
    def italic(self):
        return self._italic

    @italic.setter
    def italic(self, is_italic):
        self._italic = is_italic

    @property
    def underline(self):
        return self._underline

    @underline.setter
    def underline(self, is_underline):
        self._underline = is_underline

    @property
    def size(self):
        return self._size

    @size.setter
    def size(self, s):
        self._size = s


class TextCollection(object):
    def __init__(self, input_text, **kwargs):
        """
        Use this class when inline formatting is required in paragraphs.
        For example:
         "Only one word should be <b>bold</b> in the sentence."
         can be created as:
         TextCollection(["Only one word should be ",bold(bold)," in this sentence."])


        :param input_text:
            The input can be basestring, Text, TextCollection or a tuple of these,
                but after creation it will be a tuple of Texts.
        :param kwargs:
            bold : bool
                Bold texts.
            italic : bool
                Italic text.
            underline : bool
                Underlined texts.
            size : int
                Size of texts
        """

        self.texts = []
        # It might receive a Paragraph as text (reportlab - compatibility)
        if isinstance(input_text, Paragraph):
            input_text = input_text._text_collection.texts
        # _build_text_from_tuple requests a tuple, so if it is basestring or Text we create a tuple from it
        if not isinstance(input_text, (tuple, list, set)) or isinstance(input_text, basestring):
            input_text = [input_text]

        self.texts.extend(self._build_text_from_tuple(input_text, **kwargs))
        return

    def __add__(self, other):
        # Adding it to an other 'basestring', 'Text' or 'TextCollection' (or to a tuple of these)
        # A new TextCollection will be made in order to keep the different styles
        return TextCollection([self, other])

    def __radd__(self, other):
        # Adding it to an other 'basestring', 'Text' or 'TextCollection' (or to a tuple of these)
        # A new TextCollection will be made in order to keep the different styles
        return TextCollection([other, self])

    def add_to_paragraph(self, paragraph):
        for t in self.texts:
            t.add_to_paragraph(paragraph)
        return

    def get_texts(self, **kwargs):
        self.texts = [Text(text, **kwargs) for text in self.texts]
        return self.texts

    def _build_text_from_tuple(self, input_text, **kwargs):
        temp_texts = []
        for t in input_text:
            if isinstance(t, basestring):
                # simple string for backward compatibility, create Text with formatting and append to TextCollection
                temp_texts.append(Text(t, **kwargs))
            elif isinstance(t, Text):
                # simple Text, reformat if needed
                t = Text(t, **kwargs)
                temp_texts.append(t)
            elif isinstance(t, TextCollection):
                # The get_texts function handles reformatting, then extend the TextCollection
                temp_texts.extend(t.get_texts(**kwargs))
            elif isinstance(t, collections.Iterable) and not isinstance(t, basestring):
                # If it is a tuple then unwrap it with calling this function again
                temp_texts.extend(self._build_text_from_tuple(t, **kwargs))
        return temp_texts


class Paragraph(object):
    def __init__(self, input_text, style=None, **kwargs):
        """
        Create Paragraph

        :param input_text:
            The text of the paragraph. A TextCollection variable will be created from it.
            It can be:
                -basestring
                -Text
                -Tuple of basestrings and Texts (recursively)
                -TextCollection, the same as Tuples but during List creation
                    only this can be used for inline formatting
                -Paragraph
        :param style:
            Style of the paragraph.
                for listing all style call:
                    from docx import Document
                    d = Document()
                    styles = d.styles
                    for s in styles:
                        print s.name
        :param kwargs:
            bold : bool
                Bold texts.
            italic : bool
                Italic text.
            underline : bool
                Underlined texts.
            size : int
                Size of texts
            alignment : string ['LEFT', 'CENTER', 'RIGHT']
                Alignment of paragraph
        """

        if isinstance(input_text, Paragraph):
            # inherit
            self._style = input_text._style
            self._alignment = input_text._alignment
            # Get overrides
            if style is not None:
                self._style = style
            alignment = kwargs.pop('alignment', None)
            if alignment is not None:
                self._alignment = alignment
        else:
            if style is None:
                style = 'Normal'
            self._style = style
            self._alignment = kwargs.pop('alignment', 'LEFT')
        self._text_collection = TextCollection(input_text, **kwargs)
        return

    @property
    def style(self):
        return self._style

    def add_to_document(self, document):
        # create the paragraph or heading
        d = {'Heading1': 'Heading1', 'Heading2': 'Heading2', 'Heading3': 'Heading3', 'Normal': 'CNormal'}
        p_temp = document.add_paragraph('', style=d[self._style])
        if self._alignment is not None:
            d = {'LEFT': WD_PARAGRAPH_ALIGNMENT.LEFT, 'CENTER': WD_PARAGRAPH_ALIGNMENT.CENTER,
                 'RIGHT': WD_PARAGRAPH_ALIGNMENT.RIGHT}
            p_temp.alignment = d[self._alignment]

        # Fill paragraph with text
        self._text_collection.add_to_paragraph(p_temp)

    def add_to_paragraph(self, p_temp):
        # Fill paragraph with text
        p_temp.style = self._style
        self._text_collection.add_to_paragraph(p_temp)


class IndexedParagraph(Paragraph):
    def __init__(self, input_text, style="Normal", **kwargs):
        super(IndexedParagraph, self).__init__(input_text, style, **kwargs)

        # TODO add bookmark feature
        """
        < w:p >
            < w:pPr >
                < w:pStyle w:val = "Heading1" / >
            < / w:pPr >
            < w:bookmarkStart  w:id = "1"   w:name = "_Toc236597049" / >
            < w:r >
                < w:t > H2 < / w:t >
            < / w:r >
            < w:bookmarkEnd w:id = "1" / >
        < / w:p >
        """


class TableOfContent(object):
    toclevels = 3

    def __init__(self):
        self._toc_styles = {1: 'Heading1', 2: 'Heading2', 3: 'Heading3'}

    def add_to_document(self, document):
        paragraph = document.add_paragraph('', style='CNormal')

        fld_char = OxmlElement('w:fldChar')  # creates a new element
        fld_char.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        # fldChar.set(qn("w:dirty"), "true")
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')  # sets attribute on element
        headings = []
        if TableOfContent.toclevels > 3:
            TableOfContent.toclevels = 3
        for i in range(1, TableOfContent.toclevels+1):
            headings += [self._toc_styles[i]+";"+str(i)]
        if len(headings) == 0:
            return
        else:
            string = ';'.join(headings)
        instr_text.text = 'TOC \\t "'+string+'\h \u"'
        # 'TOC \o "1-3" \h \z \u'   # change 1-3 depending on heading levels you need

        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'separate')
        fld_char3 = OxmlElement('w:t')
        fld_char3.text = "Right-click to update field."
        fld_char2.append(fld_char3)

        fld_char4 = OxmlElement('w:fldChar')
        fld_char4.set(qn('w:fldCharType'), 'end')

        run = paragraph.add_run()
        r_element = run._r
        r_element.append(fld_char)
        r_element.append(instr_text)
        r_element.append(fld_char2)
        r_element.append(fld_char4)


class Table(object):
    def __init__(self, data, style=None, **kwargs):
        """
        Create Table
            Table( data, style,
        Note: formatting comes from Reportlab
            Where formatting commands are used always on a box of cells
                Box  : Can be defined with the left upper and right bottom cell
                Cell : Can be defined with the column and row indices.
                       WHERE THE FIRST IS THE COLUMN  - [i,j] means ith column and jth row.
                       i and j can be negative like in the python language
                       e.g.:  [5,2]   - cell of 5th column and 2nd row
                              [-1,-1] - right bottom cell

        :param data:
            The data of the Table. Which should be a tuple of rows,
            hence the first index gives back the row the second the cell in it(column)
        :param style:
            style:
                Built in office styles.
                e.g.: Table Grid, MessageTable, Light Shading Accent 1, Medium Shading 2 Accent 6, etc.
                e.g: Table(data,style=('MessageTable'))

            SPAN:
                Merging cells: ('SPAN',luc, rbc)
                    luc    : left upper corner -(col,row)
                    rbc    : left upper corner -(col,row)
                    e.g.: ('SPAN',(0,1),(1,2))

            BORDER COMMANDS:
            GRID, BOX, LINEABOVE, LINEBELOW, LINEBEFORE, LINEAFTER:
                Due to performance issues use LINEAFTER and LINEBELOW if possible.

                Creating border around cells: (<Border Command>, luc, rbc, width, color, None, style1, None, style2)
                Commands with * must be given every time. (None values are there due to compatibility reasons)
                    luc*    : left upper corner -(col,row)
                    rbc*    : left upper corner -(col,row)
                    width*  : width of line - default value is 0.25 (in office it is 4, so in the code we multiply by 8)
                    color*  : color in rgb -(float,float,float) or (255,255,255)
                    style1  : repeating pattern to create dotted or dashed line
                              single  : None
                              dotted  : (1,1)
                              dashed  : (2,2)
                              dotDash : (1,2)
                    style2  : number of parallel lines (style1 and style2 cannot be used together!)
                              single : 1
                              double : 2
                              triple : 3
                    e.g.: ('GRID',(0,0),(-1,-1),(0,0,0)) - Create Black grid
                          ('LINEBELOW',(0,0),(-1,0),(0,0,1.0), None, (2,2))
                             - Create blue dashed line below the first row
                          ('LINEAFTER',(0,0),(0,-1),(0,0,0), None, None,None, 2)
                             - Create black double line after the first column

            ALIGN:
                Align cells horizontally in box: ('ALIGN', luc, rbc, direction)
                    luc   : left upper corner -(col,row)
                    rbc   : right bottom corner -(col,row)
                    direction : ['LEFT', 'CENTER', 'RIGHT']

            VALIGN:
                Align cells vertically in box: ('ALIGN', luc, rbc, direction)
                    luc   : left upper corner -(col,row)
                    rbc   : right bottom corner -(col,row)
                    direction : ['BOTTOM', 'MIDDLE', 'TOP']

            TEXT_STYLE:
               Text style in box: ('TEXT_STYLE',luc, rbc, **styles)
                    luc   : left upper corner -(col,row)
                    rbc   : right bottom corner -(col,row)
                    **styles: ['bold', 'italic', 'underline', 'size')]
                        size : int - font size
                    e.g.: ('TEXT_STYLE', (0,0),(-1,-1),{'bold'=True,'size'=8}])

            TEXT_DIRECTION:
                Direction of vertical text: ('TEXT_DIRECTION',luc, rbc, direction)
                    luc   : left upper corner -(col,row)
                    rbc   : right bottom corner -(col,row)
                    direction : ['LEFT', 'RIGHT']
                    e.g.: ('TEXT_DIRECTION', (0,0),(-1,-1),'LEFT')


        :param **kwargs
            colWidths:
                Widths of columns: [width, width, ...]
                width : int or None

                -It always has to be the same size as the number of columns.
                -Columns that should be automatically sized have to be None.
                -Reportlab uses 3 times bigger values, thus always write 3 times bigger
                    values here to get the correct value for office
                    e.g.: [30, None, None, None] - sets the first column's width to 10 mm

            rowHeights:
                Heights of rows: [height, height, ...]
                height : int or None

                It always has to be the same size as the number of rows.
                Rows that should be automatically sized have to be None.
                    e.g.: [20, 10, None, None]

            hAlign:
                The alignment of the table on the page : direction
                direction = ['LEFT', 'CENTER', 'RIGHT']


        for listing all styles call:
            from docx import Document
            d = Document()
            styles = d.styles
            for s in styles:
                print s.name
        """

        # Store data and style
        self._data = data
        self._style = style
        self._kwargs = kwargs
        self._table = None
        self._rows = 0
        self._cols = 0
        # Initialize table size
        if isinstance(data, collections.Iterable) and not isinstance(data, basestring):
            if len(data) > 0:  # There are rows
                if isinstance(data[0], collections.Iterable) and not isinstance(data[0], basestring):  # columns
                    self.rows = len(data)
                    self.cols = len(data[0])

    def add_to_document(self, document):
        """
        :param document:
        :return:

        Merged cells can be accessed by there original index even after the merging,
        the texts assigned to these cells will be concatenated in the merged cell.

        """
        # Initialize table with size
        self._table = document.add_table(rows=self.rows, cols=self.cols)  # Initialize size
        self._table.autofit = False
        # Add a spacer() after the table
        Spacer().add_to_document(document)
        # Checking kwargs - these are coming from Reportlab
        column_widths = self._kwargs.pop('colWidths', None)
        row_heights = self._kwargs.pop('rowHeights', None)
        horizontal_align = self._kwargs.pop('hAlign', 'CENTER')

        if column_widths is not None:
            column_widths = [colWidth/3 if colWidth is not None else None for colWidth in column_widths]
            self._cols_width(column_widths)
        if row_heights is not None:
            row_heights = [rowHeight/3 if rowHeight is not None else None for rowHeight in row_heights]
            self._rows_height(row_heights)
        self._table_alignment(horizontal_align)

        # set default style for cells
        for row in self._table.rows:
            for cell in row.cells:
                p = cell.paragraphs[-1]
                p.style = document.styles['CNormal']

        # Checking styles
        if isinstance(self._style, basestring):
            self._table.style = self._style
        elif isinstance(self._style, collections.Iterable) and not isinstance(self._style, basestring):
            for s in self._style:
                if isinstance(s, collections.Iterable) and not isinstance(s, basestring):
                    print s[0]
                    if s[0] == 'GRID':
                        style = self._cell_border_style2word(s)
                        # If it should be applied for the whole table, then we use _table_border
                        if list(s[1]) == [0, 0] and list(s[2]) == [-1, -1]:
                            self._table_border(style)
                            continue

                        luc = self._get_base_indices(s[1])
                        rbc = self._get_base_indices(s[2])

                        for i in range(luc[0], rbc[0]+1):  # column indices
                            for j in range(luc[1], rbc[1]+1):  # row indices
                                sides = ['end', 'bottom']
                                if i == 0:
                                    sides += ['start']
                                if j == 0:
                                    sides += ['top']
                                self._cell_border(self._table.cell(j, i), sides, style)

                    elif s[0] == 'LINEABOVE' or s[0] == 'LINEBELOW' or s[0] == 'LINEBEFORE' or s[0] == 'LINEAFTER':
                        luc = self._get_base_indices(s[1])
                        rbc = self._get_base_indices(s[2])

                        # Always use 'end' and 'bottom', as properties above and before will overwrite 'top' and 'start'
                        # Hence we use 'top' and 'start' only if the cell is in the first row or column!
                        sides = []
                        ri = []
                        rj = []
                        if s[0] == 'LINEABOVE':
                            if luc[1] > 0:
                                sides = ['bottom']
                                ri = range(luc[0], rbc[0] + 1)  # col
                                rj = range(luc[1]-1, rbc[1])  # row
                            else:
                                sides = ['top']
                                ri = range(luc[0], rbc[0] + 1)  # col
                                rj = range(luc[1], rbc[1] + 1)  # row
                        elif s[0] == 'LINEBELOW':
                            sides = ['bottom']
                            ri = range(luc[0], rbc[0] + 1)  # col
                            rj = range(luc[1], rbc[1] + 1)  # row
                        elif s[0] == 'LINEBEFORE':
                            if luc[0] > 0:
                                sides = ['end']
                                ri = range(luc[0] - 1, rbc[0])  # col
                                rj = range(luc[1], rbc[1] + 1)  # row
                            else:
                                sides = ['start']
                                ri = range(luc[0], rbc[0] + 1)  # col
                                rj = range(luc[1], rbc[1] + 1)  # row
                        elif s[0] == 'LINEAFTER':
                            sides = ['end']
                            ri = range(luc[0], rbc[0] + 1)  # col
                            rj = range(luc[1], rbc[1] + 1)  # row
                        style = self._cell_border_style2word(s)
                        # range goes for rbc[*]-1 but we need rbc[*] as well
                        for i in ri:  # column indices
                            for j in rj:  # row indices
                                self._cell_border(self._table.cell(j, i), sides, style)

                    elif s[0] == 'BOX':
                        luc = self._get_base_indices(s[1])
                        rbc = self._get_base_indices(s[2])
                        style = self._cell_border_style2word(s)
                        # Left and Right Column
                        for j in range(luc[1], rbc[1]+1):  # row indices
                            if luc[0] == 0:
                                self._cell_border(self._table.cell(j, 0), ['start'], style)
                            else:
                                self._cell_border(self._table.cell(j, luc[0]-1), ['end'], style)
                            self._cell_border(self._table.cell(j, rbc[0]), ['end'], style)
                        # Top and Bottom cells
                        for i in range(luc[0], rbc[0]+1):  # row indices
                            if luc[1] == 0:
                                self._cell_border(self._table.cell(0, i), ['top'], style)
                            else:
                                self._cell_border(self._table.cell(luc[1]-1, i), ['bottom'], style)
                            self._cell_border(self._table.cell(rbc[1], i), ['bottom'], style)

                    elif s[0] == 'INNERGRID':
                        luc = self._get_base_indices(s[1])
                        rbc = self._get_base_indices(s[2])
                        style = self._cell_border_style2word(s)
                        for i in range(luc[0], rbc[0]+1):  # column indices
                            for j in range(luc[1], rbc[1]+1):  # row indices
                                sides = []
                                if i < rbc[0]:
                                    sides += ['end']
                                if j < rbc[1]:
                                    sides += ['bottom']
                                self._cell_border(self._table.cell(j, i), sides, style)

                    elif s[0] == 'SPAN' and len(s) == 3:
                        luc = self._get_base_indices(s[1])
                        rbc = self._get_base_indices(s[2])
                        self._span(luc, rbc)

                    elif s[0] == 'ALIGN' and len(s) == 4:
                        luc = self._get_base_indices(s[1])
                        rbc = self._get_base_indices(s[2])
                        for i in range(luc[0], rbc[0]+1):  # column indices
                            for j in range(luc[1], rbc[1]+1):  # row indices
                                self._cell_alignment(self._table.cell(j, i), s[3])

                    elif s[0] == 'VALIGN' and len(s) == 4:
                        luc = self._get_base_indices(s[1])
                        rbc = self._get_base_indices(s[2])
                        for i in range(luc[0], rbc[0]+1):  # column indices
                            for j in range(luc[1], rbc[1]+1):  # row indices
                                self._cell_valignment(self._table.cell(j, i), s[3])

                    elif s[0] == 'TEXT_STYLE' and len(s) == 4:
                        if not isinstance(s[3], dict) or isinstance(s, basestring):
                            raise TypeError("Styles must be an instance of Dictionary even for only one, not {}".format(
                                type(s[1])))
                        luc = self._get_base_indices(s[1])
                        rbc = self._get_base_indices(s[2])

                        for i in range(luc[0], rbc[0]+1):  # column indices
                            for j in range(luc[1], rbc[1]+1):  # row indices
                                self._cell_text_style((j, i), s[3])

                    elif s[0] == 'TEXT_DIRECTION' and len(s) == 4:
                        luc = self._get_base_indices(s[1])
                        rbc = self._get_base_indices(s[2])

                        for i in range(luc[0], rbc[0]+1):  # column indices
                            for j in range(luc[1], rbc[1]+1):  # row indices
                                self._cell_direction(self._table.cell(j, i), s[3])

                elif isinstance(s, basestring):
                    self._table.style = s

        # Fill table with text
        for r, row in zip(range(len(self._table.rows)), self._table.rows):
            for c, cell in zip(range(len(row.cells)), row.cells):
                p = cell.paragraphs[-1]  # There is already a default paragraph in every cell
                if isinstance(self._data[r][c], basestring):
                    Text(self._data[r][c]).add_to_paragraph(p)
                elif isinstance(self._data[r][c], (Text, TextCollection, Paragraph)):
                    self._data[r][c].add_to_paragraph(p)
                elif isinstance(self._data[r][c], (tuple, list, set)):
                    self._add_tuple_text_to_cell(cell, p, self._data[r][c])

    def _add_tuple_text_to_cell(self, cell, p, data):
        if len(data) == 0:
            return
        i = 0
        i_max = len(data)
        for t in data:
            if isinstance(t, basestring):
                Text(t).add_to_paragraph(p)
                if i < i_max-1:
                    alignment = p.alignment
                    p = cell.add_paragraph('', style=p.style)  # we inherit the previous style
                    p.alignment = alignment

            elif isinstance(t, (Text, TextCollection, Paragraph)):
                t.add_to_paragraph(p)
                if i < i_max-1:
                    alignment = p.alignment
                    p = cell.add_paragraph('', style=p.style)  # we inherit the previous style
                    p.alignment = alignment
            if isinstance(t, (tuple, list, set)):
                self._add_tuple_text_to_cell(cell, p, t)
            i += 1

    def _span(self, luc, rbc):
        luc = self._get_base_indices(luc)
        rbc = self._get_base_indices(rbc)
        lu_cell = self._table.cell(luc[1], luc[0])
        rb_cell = self._table.cell(rbc[1], rbc[0])
        lu_cell.merge(rb_cell)

    def _get_base_indices(self, c):
        if not isinstance(c, (list, tuple, set)):
            raise TypeError("Indices must be an (int,int), not '{}'".format(type(c)))
        cell = list(c)
        cell[0] = self.cols + c[0] if c[0] < 0 else c[0]
        cell[1] = self.rows + c[1] if c[1] < 0 else c[1]
        return cell

    @staticmethod
    def _cell_border_style2word(s):
        # Create styles
        # Width
        if not isinstance(s[3], (int, float)):
            raise TypeError("Grid line width should be int or float, not '{}'".format(type(s[3])))
        # Default width is 0.25 for Reportlab, it is 4 for Word, hence multiply it with 8
        width = str(int(s[3]*8))
        # Color
        if not isinstance(s[4], collections.Iterable):
            raise TypeError("Color of the lines should be (float,float,float), not '{}'".format(type(s[4])))
        color = ''
        for c in s[4]:
            if not isinstance(c, float):
                raise TypeError(
                    "Color of the lines should be float, not '{}'".format(type(c)))
            color += '{0:02X}'.format(int(c * 255)) if isinstance(c, float) else '{0:02X}'.format(c)
        # Line style. In office there is no dotted double, hence double and triple has priority over pattern
        line_style = 'single'
        if len(s) > 6:
            if s[6] is None:
                line_style = 'single'
            elif list(s[6]) == list((1, 1)):  # dot style
                line_style = 'dotted'
            elif list(s[6]) == list((2, 2)):
                line_style = 'dashed'
            elif list(s[6]) == list((1, 2)):
                line_style = 'dotDash'
        if len(s) > 8:  # There could be even more, but those are for Reportlab
            if s[8] == 2:
                line_style = 'double'
            elif s[8] == 3:
                line_style = 'triple'

        style = {'width': width, 'color': color, 'style': line_style}
        return style

    @staticmethod
    def _cell_border(cell, sides, style):
        tc_pr_xml = cell._tc.get_or_add_tcPr()

        border = tc_pr_xml.find(qn('w:tcBorders'))
        if border is None:
            border = OxmlElement('w:tcBorders')  # then create new
            for side in sides:
                line = OxmlElement('w:{}'.format(side))
                line.set(qn('w:val'), '{}'.format(style['style']))
                line.set(qn('w:sz'), '{}'.format(style['width']))
                line.set(qn('w:color'), '{}'.format(style['color']))
                line.set(qn('w:space'), '0')
                border.append(line)
        else:
            for side in sides:
                line = border.find(qn('w:{}'.format(side)))
                if line is not None:
                    border.remove(line)
                line = OxmlElement('w:{}'.format(side))
                line.set(qn('w:val'), '{}'.format(style['style']))
                line.set(qn('w:sz'), '{}'.format(style['width']))
                line.set(qn('w:color'), '{}'.format(style['color']))
                line.set(qn('w:space'), '0')
                border.append(line)
            tc_pr_xml.remove(border)
        tc_pr_xml.append(border)

    def _table_border(self, style):
        sides = ['insideH', 'insideV', 'top', 'bottom', 'left', 'right']
        tbl_pr_xml = self._table._tblPr
        border = tbl_pr_xml.find(qn('w:tblBorders'))
        if border is None:
            border = OxmlElement('w:tblBorders')
            for side in sides:
                line = OxmlElement('w:{}'.format(side))
                line.set(qn('w:val'), '{}'.format(style['style']))
                line.set(qn('w:sz'), '{}'.format(style['width']))
                line.set(qn('w:color'), '{}'.format(style['color']))
                line.set(qn('w:space'), '0')
                border.append(line)
        else:
            for side in sides:
                line = border.find(qn('w:{}'.format(side)))
                if line is not None:
                    border.remove(line)
                line = OxmlElement('w:{}'.format(side))
                line.set(qn('w:val'), '{}'.format(style['style']))
                line.set(qn('w:sz'), '{}'.format(style['width']))
                line.set(qn('w:color'), '{}'.format(style['color']))
                line.set(qn('w:space'), '0')
                border.append(line)
            tbl_pr_xml.remove(border)
        tbl_pr_xml.append(border)

    @staticmethod
    def _cell_direction(cell, orientation):
        # Get header in xml
        tc_pr_xml = cell._tc.get_or_add_tcPr()
        # Create border segment
        text_direction = OxmlElement('w:textDirection')
        d = {'RIGHT': 'tbRl', 'LEFT': 'btLr'}
        text_direction.set(qn('w:val'), '{}'.format(d[orientation]))
        tc_pr_xml.append(text_direction)

    @staticmethod
    def _cell_alignment(cell, direction):
        for p in cell.paragraphs:
            d = {'LEFT': WD_PARAGRAPH_ALIGNMENT.LEFT, 'CENTER': WD_PARAGRAPH_ALIGNMENT.CENTER,
                 'RIGHT': WD_PARAGRAPH_ALIGNMENT.RIGHT}
            p.alignment = d[direction]

    @staticmethod
    def _cell_valignment(cell, direction):
        # Get header in xml
        tc_pr_xml = cell._tc.get_or_add_tcPr()
        vertical_align = tc_pr_xml.find(qn('w:vAlign'))
        if vertical_align is not None:
            tc_pr_xml.remove(vertical_align)
        # Create border segment
        vertical_align = OxmlElement('w:vAlign')
        d = {'BOTTOM': 'bottom', 'MIDDLE': 'center', 'TOP': 'top'}  # Reportlab to docx convert
        vertical_align.set(qn('w:val'), '{}'.format(d[direction]))
        tc_pr_xml.append(vertical_align)

    def _cell_text_style(self, c, styles):
        # c[0] is the row and c[1] is the column, it is already switched !!
        self._data[c[0]][c[1]] = decorate(self._data[c[0]][c[1]], **styles)

    def _cols_width(self, widths):
        if isinstance(widths, (list, tuple, set)):
            if len(widths) != self.cols:
                raise TypeError("Not enough number of column widths. It is :{} but it should be: {}".format(
                    len(widths), self.cols))
        else:
            raise TypeError("Widths should be an Iterable, not '{}'".format(type(widths)))
        for i in range(self.cols):
            if widths[i] is None:
                continue  # Do not want to change this
            for cell in self._table.column_cells(i):
                if isinstance(widths[i], int):
                    cell.width = Mm(widths[i])
                else:
                    raise TypeError("Width should be an int, not '{}'".format(type(widths[i])))

    def _rows_height(self, heights):
        if isinstance(heights, (list, tuple, set)):
            if len(heights) != self.rows:
                raise TypeError("Not enough number of row heights. It is :{} but it should be: {}".format(
                    len(heights), self.rows))
        else:
            raise TypeError("Widths should be an Iterable, not '{}'".format(type(heights)))
        for i in range(self.rows):
            # Get row
            if heights[i] is None:  # Do not want to change this, Reportlab compatibility
                continue
            row = self._table.rows[i]
            if isinstance(heights[i], int):
                h = heights[i]
            else:
                raise TypeError("Height should be an instance of int, not {}".format(type(heights)))
            # Get header in xml
            tr_pr_xml = row._tr.get_or_add_trPr()  # No built in property in python-docx

            row_height = tr_pr_xml.find(qn('w:trHeight'))
            if row_height is not None:
                tr_pr_xml.remove(row_height)
            # Create Height segment
            row_height = OxmlElement('w:trHeight')
            row_height.set(qn('w:val'), str(h) + 'mm')
            row_height.set(qn('w:hRule'), 'exact')
            tr_pr_xml.append(row_height)

    def _table_alignment(self, direction):
        d = {'LEFT': WD_TABLE_ALIGNMENT.LEFT, 'CENTER': WD_TABLE_ALIGNMENT.CENTER,
             'RIGHT': WD_TABLE_ALIGNMENT.RIGHT}
        self._table.alignment = d[direction]


class NonEmptyTable(Table):
    def __init__(self, data, style, **kwargs):
        super(NonEmptyTable, self).__init__(data=data, style=style, **kwargs)

    def add_to_document(self, document):
        if self.rows > 0:
            super(NonEmptyTable, self).add_to_document(document)
        else:
            Paragraph("No items to display.").add_to_document(document)


class NonEmptyTableWithHeader(Table):
    def __init__(self, data, style, **kwargs):
        super(NonEmptyTableWithHeader, self).__init__(data=data, style=style, **kwargs)

    def add_to_document(self, document):
        if self.rows > 1:
            super(NonEmptyTableWithHeader, self).add_to_document(document)
        else:
            Paragraph("No items to display.").add_to_document(document)


class List:
    def __init__(self, texts, style='CNormal', init_level=0, **kwargs):
        """
        Create List

        :param texts:
            The texts can be string, Text or TextCollection. Do not use the tuple method for creating Texts with
            inline formatting as here it will be interpreted as different levels. For inline formatting use
            TextCollections
            e.g.: List([1,['a'],[[TextCollection(['This ',bold('is'),' a text'])]]], style)

                * 1
                    * a
                        * This is a text
            while: List([1,['a'],[['This ',bold('is'),' a text']]], style)
                * 1
                    * a
                        * This
                        * is
                        * a text
        """

        self.texts = texts
        self.symbol = kwargs.pop('symbol', 'bullet')
        # TODO change hardcoded CNormal
        self.style = 'CNormal'
        self._document = None
        self._name = None
        self._paragraph_styles = None
        self._numbering_id = None

        # TODO add new level definitions
        self._max_level = 3

    def add_to_document(self, document):
        self._document = document
        # Create new abstract based on templates.py
        from doctemplate import template_loader
        self._paragraph_styles, self._numbering_id = template_loader.create_new_numbering_style(self.symbol)

        # If it is only a string or Text
        if isinstance(self.texts, basestring):
            p = document.add_paragraph('', style=self.style)
            Text(self.texts).add_to_paragraph(p)
        elif isinstance(self.texts, Text) or isinstance(self.texts, TextCollection):
            p = document.add_paragraph('', style=self.style)
            self.texts.add_to_paragraph(p)
        # If it is an actual List
        elif isinstance(self.texts, collections.Iterable) and not isinstance(self.texts, basestring):
            self._build_list_from_tuple(self.texts)

    def _build_list_from_tuple(self, texts, level=0):
        if level > self._max_level-1:
            level = self._max_level-1

        for text in texts:
            # Get the text or unwrap the next level
            if isinstance(text, basestring):  # simple string for backward compatibility
                p = self._document.add_paragraph('', style=self._paragraph_styles+str(level+1))
                Text(text).add_to_paragraph(p)
            elif isinstance(text, Text) or isinstance(text, TextCollection):
                p = self._document.add_paragraph('', style=self._paragraph_styles+str(level+1))
                text.add_to_paragraph(p)
            elif isinstance(text, collections.Iterable) and not isinstance(text, basestring):
                self._build_list_from_tuple(text, level=level+1)


class Image:
    def __init__(self, path, width=None, height=None, **kwargs):
        """
        Adding image to the docx
        :param
            path:
                path to the image
            width:
                Desired width of image. Default is the original size.
            height:
                Desired height of image. If this is not assigned just the width then the picture
                is sized proportionally
        :Keywords:
            hAlign : string
                Image alignment: ALIGN=direction
                direction: ['LEFT', 'CENTER', 'RIGHT']
                e.g.: hAlign='RIGHT')
        """
        self.path = path
        self.align = kwargs.pop('hAlign', '')
        self.width = width
        self.height = height
        return

    def add_to_document(self, document):
        if self.width is not None:
            if self.height is not None:
                document.add_picture(self.path, self.width, self.height)
            else:
                document.add_picture(self.path, self.width)
        else:
            document.add_picture(self.path)

        pic_paragraph = document.paragraphs[-1]
        if self.align == 'LEFT':
            pic_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif self.align == 'CENTER':
            pic_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif self.align == 'RIGHT':
            pic_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


class Header(object):
    # TODO create xml code - Implementation is in process in pyhton-docx
    def __init__(self):
        self._story = []

    def add_element(self, element):
        self._story += [element]

    def _add_to_section(self, section):
        pass

    def add_to_document(self, document):
        pass
        # header = document.sections[0].header
        # self._add_to_section(header)


def bold(text, is_it=True):
    if isinstance(text, Text) or isinstance(text, basestring) or isinstance(text, TextCollection):
        text = TextCollection(text, bold=is_it)
    else:
        raise TypeError("The text must be an instance of 'Text',"
                        " 'TextCollection' or 'basestring', not '{}'".format(type(text)))
    return text


def underline(text, is_it=True):
    if isinstance(text, Text) or isinstance(text, basestring) or isinstance(text, TextCollection):
        text = TextCollection(text, underline=is_it)
    else:
        raise TypeError("The text must be an instance of 'Text',"
                        " 'TextCollection' or 'basestring', not '{}'".format(type(text)))
    return text


def italic(text, is_it=True):
    if isinstance(text, Text) or isinstance(text, basestring) or isinstance(text, TextCollection):
        text = TextCollection(text, italic=is_it)
    else:
        raise TypeError("The text must be an instance of 'Text',"
                        " 'TextCollection' or 'basestring', not '{}'".format(type(text)))
    return text


def size(text, s):
    if not isinstance(s, int):
        raise TypeError("Text size must be an int, not '{}'".format(type(s)))
    if isinstance(text, Text) or isinstance(text, basestring) or isinstance(text, TextCollection):
        text = TextCollection(text, size=s)
    else:
        raise TypeError("The text must be an instance of 'Text',"
                        " 'TextCollection' or 'basestring', not '{}'".format(type(text)))
    return text


def decorate(text, **kwargs):
    """
    Format text, Text, TextCollection or Paragraph.

    :Keywords:
      bold : bool
        Bold text.
      italic : bool
        Italic text.
      underline : bool
        Underlined text.
      size : int
        The size of the text

    :ReturnType: str
    :Return: Formatted text.
    """
    if isinstance(text, (list, tuple, set)):   # In order to add standalone Paragraphs
        return [decorate(t, **kwargs) for t in text]
    elif isinstance(text, Text) or isinstance(text, basestring) or isinstance(text, TextCollection):
        text = TextCollection(text, **kwargs)
    elif isinstance(text, Paragraph):
        text = Paragraph(text, **kwargs)
    else:
        raise TypeError("The text of decorate must be an instance of 'Text', 'basestring', "
                        "'Paragraph' or TextCollection or a tuple of these, not '{}'".format(type(text)))
    return text
