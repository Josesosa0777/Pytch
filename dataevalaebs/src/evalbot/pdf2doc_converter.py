import argparse
import os
import re
import docx
import mammoth
from pdf2docx import Converter


def add_headings_to_doc(docx_file_path, dest_file_path=None):
    file_path, ext = os.path.splitext(docx_file_path)
    regex_List = ['.* sec$',
                  '^\d*. Overall summary$',
                  '^\d. .* event details$',
                  '^\d*. Lane \w.*$','^\d*. Blockage \w.*$'
                  ]

    if ext not in [".docx", ".doc"]:
        print("Please provide valid doc/docx input file")
        return False

    try:
        doc = docx.Document((docx_file_path))

        # get all paragraphs
        all_paras = doc.paragraphs

        # match the expression and add heading in doc
        flag = 0

        for paragraph in all_paras:
            result = re.compile('|'.join(regex_List)).findall(
                paragraph.text)

            if len(result) != 0:
                for res in result:

                    if ('' + res).endswith("Overall summary"):
                        flag = 1
                    if flag == 1:
                        if (''+res).endswith('Overall summary') or (''+res).endswith('event details') or (''+res).__contains__('Lane') or (''+res).__contains__('Blockage'):#if re.match('^\d*. .*$', res) is not None:
                            paragraph.style = doc.styles['Heading 1']
                            recent_res=res
                        else:
                            if (''+recent_res).endswith('Overall summary'):
                                paragraph.style = doc.styles['Heading 1']
                            else:
                                paragraph.style = doc.styles['Heading 2']
        if dest_file_path is not None:
            dest_docx_file_path = dest_file_path
        else:
            dest_docx_file_path = docx_file_path

        doc.save(dest_docx_file_path)
        print("Source File: {}\nDestination Path: {}".format(docx_file_path, dest_docx_file_path))
        return True
    except Exception as e:
        print(str(e))
        return False


def convert_pdf_2_docx(pdf_file_path, docx_file_path):
    # Check files
    _, ext = os.path.splitext(pdf_file_path)
    if ext != '.pdf':
        print("File extension should be '.pdf', not '{}'".format(ext))
        return False
    _, ext = os.path.splitext(docx_file_path)
    if ext == '.doc':
        docx_file_path = docx_file_path.replace(ext, '.docx')
    elif ext != '.docx':
        print("File extension should be '.docx', not '{}'".format(ext))
        docx_file_path = docx_file_path.replace(ext, '.docx')
        print(docx_file_path)
    # convert pdf to docx
    cv = Converter(pdf_file_path)
    cv.convert(docx_file_path, start=0, end=None)
    cv.close()

    status = add_headings_to_doc(docx_file_path)
    if status:
        print("Headings added successfully.")
    else:
        print("Exception while adding headings to file")
    return docx_file_path


def docx_to_html(docx_file_path, html_file_path):
    try:
        docx_fp = open(docx_file_path, 'rb')
        html_fp = open(html_file_path, 'wb')
        document = mammoth.convert_to_html(docx_fp)
        html_fp.write(document.value.encode('utf8'))
        docx_fp.close()
        html_fp.close()
        return html_file_path
    except:
        return False


def main():
    # parse arguments
    parser = argparse.ArgumentParser(description=
                                     "Converter: Pdf to doc")
    parser.add_argument('pdf_docx',
                        nargs='+',
                        help='Input pdf file path and output docx file path')
    parser.add_argument('-html',
                        action='store_true',
                        default=False,
                        help='Do not wait for measurement files, run immediately on existing ones')
    args = parser.parse_args().pdf_docx
    html_conversion_needed = parser.parse_args().html
    argument_size = len(args)
    if argument_size > 0:
        if argument_size > 1:
            pdf_file_path = args[0]
            docx_file_path = args[1]
        else:
            pdf_file_path = args[0]
            file_path, ext = os.path.splitext(pdf_file_path)
            docx_file_path = file_path + ".docx"

        if os.path.isfile(pdf_file_path):

            status = convert_pdf_2_docx(pdf_file_path, docx_file_path)
            if status:
                print("Output doc file: {}".format(status))
                if html_conversion_needed:
                    file_path, ext = os.path.splitext(docx_file_path)
                    html_file_path = file_path + ".html"
                    result = docx_to_html(status, html_file_path)
                    print("Output html file: {}".format(result))
        else:
            print("Please provide valid pdf file path")
    else:
        print("Please provide pdf and docx file paths")
main()
